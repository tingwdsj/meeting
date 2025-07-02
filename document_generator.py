"""
文档生成模块 - 会议纪要生成神器
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class DocumentGenerator:
    """文档生成类"""
    
    def __init__(self):
        self.default_font = "微软雅黑"
        self.title_font_size = 18
        self.heading_font_size = 14
        self.body_font_size = 12
        
    def create_meeting_minutes_doc(self, 
                                  meeting_info: str, 
                                  transcription: str, 
                                  minutes: str,
                                  output_path: Optional[str] = None) -> str:
        """
        生成会议纪要Word文档
        
        Args:
            meeting_info: 会议描述信息
            transcription: 会议录音文本
            minutes: 会议纪要
            output_path: 输出文件路径
            
        Returns:
            生成的文档路径
        """
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"会议纪要_{timestamp}.docx"
        
        try:
            # 创建文档
            doc = Document()
            
            # 设置文档样式
            self._setup_document_style(doc)
            
            # 添加标题
            title = doc.add_heading("会议纪要", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加会议基本信息
            doc.add_heading("会议基本信息", level=1)
            info_para = doc.add_paragraph(meeting_info)
            info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加会议纪要
            doc.add_heading("会议纪要", level=1)
            minutes_para = doc.add_paragraph(minutes)
            minutes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加生成时间
            doc.add_paragraph()
            time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 保存文档
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"生成会议纪要文档失败: {str(e)}")
    
    def create_complete_info_doc(self, 
                                meeting_info: str, 
                                transcription: str, 
                                minutes: str,
                                output_path: Optional[str] = None) -> str:
        """
        生成会议完整信息Word文档
        
        Args:
            meeting_info: 会议描述信息
            transcription: 会议录音文本
            minutes: 会议纪要
            output_path: 输出文件路径
            
        Returns:
            生成的文档路径
        """
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"会议完整信息_{timestamp}.docx"
        
        try:
            # 创建文档
            doc = Document()
            
            # 设置文档样式
            self._setup_document_style(doc)
            
            # 添加标题
            title = doc.add_heading("会议完整信息", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加会议基本信息
            doc.add_heading("会议基本信息", level=1)
            info_para = doc.add_paragraph(meeting_info)
            info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加会议纪要
            doc.add_heading("会议纪要", level=1)
            minutes_para = doc.add_paragraph(minutes)
            minutes_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加会议录音文本
            doc.add_heading("会议录音文本", level=1)
            transcription_para = doc.add_paragraph(transcription)
            transcription_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加生成时间
            doc.add_paragraph()
            time_para = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 保存文档
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise Exception(f"生成会议完整信息文档失败: {str(e)}")
    
    def _setup_document_style(self, doc: Document):
        """
        设置文档样式
        
        Args:
            doc: Word文档对象
        """
        # 设置默认字体
        doc.styles['Normal'].font.name = self.default_font
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        doc.styles['Normal'].font.size = Pt(self.body_font_size)
        
        # 设置标题样式
        for i in range(1, 10):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                doc.styles[style_name].font.name = self.default_font
                doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                if i == 1:
                    doc.styles[style_name].font.size = Pt(self.heading_font_size)
                else:
                    doc.styles[style_name].font.size = Pt(self.heading_font_size - (i-1))
    
    def format_text_for_doc(self, text: str) -> str:
        """
        格式化文本以适应Word文档
        
        Args:
            text: 原始文本
            
        Returns:
            格式化后的文本
        """
        # 处理换行符
        text = text.replace('\n', '\n\n')
        
        # 处理特殊字符
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        return text
    
    def get_document_info(self, file_path: str) -> dict:
        """
        获取文档信息
        
        Args:
            file_path: 文档路径
            
        Returns:
            文档信息字典
        """
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            return {
                "paragraphs_count": len(paragraphs),
                "total_text_length": sum(len(p) for p in paragraphs),
                "file_size": os.path.getsize(file_path),
                "created_time": datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
            }
        except Exception as e:
            return {"error": str(e)}

# 全局文档生成器实例
document_generator = DocumentGenerator() 